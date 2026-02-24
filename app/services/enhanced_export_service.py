"""
Enhanced Export Service - Professional DOCX and PDF exports with embedded charts
NO EMOJI in exports - Professional formatting only
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
from datetime import datetime
import logging
from typing import Dict, List, Optional
import base64

try:
    import plotly.io as pio
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False
    logging.warning("Plotly not available for chart exports")

logger = logging.getLogger(__name__)


class EnhancedExportService:
    """Service for generating professional DOCX and PDF exports"""
    
    def __init__(self):
        # Configure plotly for static image export if available
        if PLOTLY_AVAILABLE:
            try:
                pio.kaleido.scope.mathjax = None
            except:
                pass
    
    def generate_comprehensive_docx(
        self,
        customer_name: str,
        profile_data: Dict,
        customer_data: Dict,
        market_intel: Dict = None,
        projects: List[Dict] = None,
        financial_data: Dict = None,
        charts: Dict = None
    ) -> BytesIO:
        """
        Generate comprehensive customer analysis DOCX
        
        Args:
            customer_name: Customer name
            profile_data: AI-generated profile
            customer_data: Raw CRM data
            market_intel: Market intelligence analysis
            projects: Project data
            financial_data: Financial analysis data
            charts: Dict of plotly figure objects {'chart_name': fig}
        
        Returns:
            BytesIO buffer containing DOCX
        """
        doc = Document()
        
        # Set up styles
        self._setup_document_styles(doc)
        
        # Title Page
        self._add_title_page(doc, customer_name)
        
        # Table of Contents (manual for docx)
        self._add_table_of_contents(doc)
        
        # Section 1: Customer Profile
        self._add_customer_profile_section(doc, profile_data, customer_data)
        
        # Section 2: Deep Dive Analytics
        if customer_data:
            self._add_deep_dive_section(doc, customer_data, charts)
        
        # Section 3: Market Intelligence
        if market_intel:
            self._add_market_intelligence_section(doc, market_intel)
        
        # Section 4: Project Analysis
        if projects:
            self._add_project_section(doc, projects, charts)
        
        # Section 5: Financial Analysis
        if financial_data:
            self._add_financial_section(doc, financial_data, charts)
        
        # Footer with generation timestamp
        self._add_footer(doc)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _setup_document_styles(self, doc: Document):
        """Configure professional document styles"""
        styles = doc.styles
        
        # Heading 1 style
        if 'Heading 1' in [s.name for s in styles]:
            h1 = styles['Heading 1']
            h1.font.name = 'Arial'
            h1.font.size = Pt(18)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(0, 51, 102)  # Professional blue
        
        # Heading 2 style
        if 'Heading 2' in [s.name for s in styles]:
            h2 = styles['Heading 2']
            h2.font.name = 'Arial'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(0, 102, 204)
        
        # Normal text
        if 'Normal' in [s.name for s in styles]:
            normal = styles['Normal']
            normal.font.name = 'Arial'
            normal.font.size = Pt(11)
    
    def _add_title_page(self, doc: Document, customer_name: str):
        """Add professional title page"""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('Customer Analysis Report')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        doc.add_paragraph()  # Spacing
        
        # Customer name
        customer_para = doc.add_paragraph()
        customer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = customer_para.add_run(customer_name)
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 102, 204)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Generation date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document):
        """Add table of contents"""
        doc.add_heading('Table of Contents', level=1)
        
        toc_items = [
            '1. Customer Profile',
            '2. Deep Dive Analytics',
            '3. Market Intelligence',
            '4. Project Analysis',
            '5. Financial Analysis'
        ]
        
        for item in toc_items:
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_customer_profile_section(self, doc: Document, profile_data: Dict, customer_data: Dict):
        """Add customer profile section"""
        doc.add_heading('1. Customer Profile', level=1)
        
        basic_data = profile_data.get('basic_data', {})
        
        # Basic Information Table
        doc.add_heading('Basic Information', level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        
        fields = [
            ('Company Name', basic_data.get('name', 'N/A')),
            ('Headquarters', basic_data.get('hq_address', 'N/A')),
            ('CEO', basic_data.get('ceo', 'N/A')),
            ('Ownership', basic_data.get('owner', 'N/A')),
            ('Employees (FTE)', basic_data.get('fte', 'N/A')),
            ('Industry', basic_data.get('company_focus', 'N/A')),
            ('Financial Status', basic_data.get('financials', 'N/A'))
        ]
        
        for label, value in fields:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
            # Bold the labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Company Overview (if available from external sources)
        if 'company_overview' in profile_data:
            doc.add_heading('Company Overview', level=2)
            overview = profile_data['company_overview']
            doc.add_paragraph(overview.get('description', 'No description available'))
            
            if overview.get('source_url'):
                p = doc.add_paragraph()
                p.add_run('Source: ')
                p.add_run(overview['source_url']).font.italic = True
        
        # Latest Developments (if available)
        if 'recent_news' in profile_data:
            doc.add_heading('Recent News & Developments', level=2)
            news_items = profile_data['recent_news']
            for idx, news in enumerate(news_items[:5], 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{news.get('title', 'No title')}").font.bold = True
                p.add_run(f" ({news.get('published_date', 'Unknown date')})")
                if news.get('description'):
                    doc.add_paragraph(news['description'], style='List Bullet 2')
                if news.get('url'):
                    p = doc.add_paragraph(style='List Bullet 2')
                    p.add_run('URL: ')
                    p.add_run(news['url']).font.italic = True
        
        doc.add_page_break()
    
    def _add_deep_dive_section(self, doc: Document, customer_data: Dict, charts: Dict):
        """Add deep dive analytics section"""
        doc.add_heading('2. Deep Dive Analytics', level=1)
        
        doc.add_heading('Key Performance Indicators', level=2)
        # KPIs would be calculated from customer_data
        doc.add_paragraph('Revenue Trends: Analysis pending')
        doc.add_paragraph('Customer Lifetime Value: Analysis pending')
        doc.add_paragraph('Engagement Score: Analysis pending')
        
        # Add charts if available
        if charts and 'revenue_trend' in charts:
            self._add_chart_to_doc(doc, charts['revenue_trend'], 'Revenue Trend Chart')
        
        doc.add_page_break()
    
    def _add_market_intelligence_section(self, doc: Document, market_intel: Dict):
        """Add market intelligence section"""
        doc.add_heading('3. Market Intelligence', level=1)
        
        sections = [
            ('Financial Health', market_intel.get('financial_health', '')),
            ('Recent Developments', market_intel.get('recent_developments', '')),
            ('Market Position', market_intel.get('market_position', '')),
            ('Strategic Outlook', market_intel.get('strategic_outlook', '')),
            ('Risk Assessment', market_intel.get('risk_assessment', ''))
        ]
        
        for title, content in sections:
            if content:
                doc.add_heading(title, level=2)
                doc.add_paragraph(content)
        
        # Competitors
        if market_intel.get('competitors'):
            doc.add_heading('Key Competitors', level=2)
            for competitor in market_intel['competitors']:
                doc.add_paragraph(competitor, style='List Bullet')
        
        # Source citations
        if market_intel.get('sources'):
            doc.add_heading('Sources', level=3)
            for source in market_intel['sources']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(source).font.italic = True
        
        doc.add_page_break()
    
    def _add_project_section(self, doc: Document, projects: List[Dict], charts: Dict):
        """Add project analysis section"""
        doc.add_heading('4. Project Analysis', level=1)
        
        doc.add_heading('Project Summary', level=2)
        doc.add_paragraph(f'Total Projects: {len(projects)}')
        
        active = sum(1 for p in projects if p.get('status') == 'Active')
        doc.add_paragraph(f'Active Projects: {active}')
        
        # Project details
        for project in projects:
            doc.add_heading(project.get('name', 'Unnamed Project'), level=3)
            
            details = [
                ('Status', project.get('status', 'Unknown')),
                ('Start Date', project.get('start_date', 'N/A')),
                ('End Date', project.get('end_date', 'N/A')),
                ('Budget', f"${project.get('budget', 0):,.2f}"),
                ('Progress', f"{project.get('progress', 0)}%")
            ]
            
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Light List Accent 1'
            
            for label, value in details:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(value)
            
            doc.add_paragraph()  # Spacing
        
        # Add Gantt chart if available
        if charts and 'gantt_chart' in charts:
            self._add_chart_to_doc(doc, charts['gantt_chart'], 'Project Timeline')
        
        doc.add_page_break()
    
    def _add_financial_section(self, doc: Document, financial_data: Dict, charts: Dict):
        """Add financial analysis section"""
        doc.add_heading('5. Financial Analysis', level=1)
        
        doc.add_heading('Cost Breakdown', level=2)
        # Cost table
        costs = financial_data.get('cost_breakdown', {})
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Medium Grid 1 Accent 1'
        
        # Header row
        header = table.add_row()
        header.cells[0].text = 'Category'
        header.cells[1].text = 'Amount'
        for cell in header.cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for category, amount in costs.items():
            row = table.add_row()
            row.cells[0].text = category
            row.cells[1].text = f"${amount:,.2f}"
        
        # Budget Variance
        if 'budget_variance' in financial_data:
            doc.add_heading('Budget Variance Analysis', level=2)
            variance = financial_data['budget_variance']
            doc.add_paragraph(f"Budgeted: ${variance.get('budgeted', 0):,.2f}")
            doc.add_paragraph(f"Actual: ${variance.get('actual', 0):,.2f}")
            doc.add_paragraph(f"Variance: ${variance.get('variance', 0):,.2f} ({variance.get('variance_percent', 0):.1f}%)")
            doc.add_paragraph(f"Status: {variance.get('status', 'Unknown')}")
        
        # Add cost trend chart if available
        if charts and 'cost_trend' in charts:
            self._add_chart_to_doc(doc, charts['cost_trend'], 'Cost Trend Analysis')
        
        doc.add_page_break()
    
    def _add_chart_to_doc(self, doc: Document, fig, caption: str):
        """Add plotly chart as image to document"""
        if not PLOTLY_AVAILABLE:
            doc.add_paragraph(f"[Chart: {caption}] - Chart export requires kaleido library")
            return
        
        try:
            # Convert plotly figure to image bytes
            img_bytes = pio.to_image(fig, format='png', width=800, height=400)
            img_stream = BytesIO(img_bytes)
            
            # Add to document
            doc.add_picture(img_stream, width=Inches(6))
            
            # Add caption
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
            
        except Exception as e:
            logger.error(f"Failed to add chart {caption}: {e}")
            doc.add_paragraph(f"[Chart: {caption}] - Error generating chart image")
    
    def _add_footer(self, doc: Document):
        """Add document footer"""
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Customer Analysis Report"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generate_filename(self, customer_name: str, extension: str) -> str:
        """Generate standardized filename"""
        safe_name = "".join(c for c in customer_name if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_name = safe_name.replace(' ', '_')
        timestamp = datetime.now().strftime('%Y%m%d')
        return f"{safe_name}_Analysis_{timestamp}.{extension}"
    
    def generate_comprehensive_pdf(
        self,
        customer_name: str,
        profile_data: Dict,
        customer_data: Dict,
        market_intel: Dict = None,
        projects: List[Dict] = None,
        financial_data: Dict = None,
        charts: Dict = None
    ) -> BytesIO:
        """
        Generate comprehensive customer analysis PDF using reportlab
        
        Args:
            customer_name: Customer name
            profile_data: AI-generated profile
            customer_data: Raw CRM data
            market_intel: Market intelligence analysis
            projects: Project data
            financial_data: Financial analysis data
            charts: Dict of plotly figure objects {'chart_name': fig}
        
        Returns:
            BytesIO buffer containing PDF
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, PageBreak,
                Table, TableStyle, Image as RLImage
            )
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        except ImportError:
            logger.error("reportlab not available for PDF generation")
            raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
        
        buffer = BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#374151'),
            spaceAfter=12,
            spaceBefore=12
        ))
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            textColor=colors.HexColor('#4b5563'),
            spaceAfter=12
        ))
        
        # Title Page
        elements.append(Spacer(1, 2*inch))
        title = Paragraph(f"<b>Customer Analysis Report</b>", styles['CustomTitle'])
        elements.append(title)
        elements.append(Spacer(1, 0.3*inch))
        
        subtitle = Paragraph(f"<b>{customer_name}</b>", styles['Heading2'])
        subtitle.alignment = TA_CENTER
        elements.append(subtitle)
        elements.append(Spacer(1, 0.2*inch))
        
        date_text = Paragraph(
            f"Generated: {datetime.now().strftime('%B %d, %Y')}",
            styles['Normal']
        )
        date_text.alignment = TA_CENTER
        elements.append(date_text)
        elements.append(PageBreak())
        
        # Section 1: Customer Profile
        elements.append(Paragraph("<b>1. Customer Profile</b>", styles['CustomHeading']))

        def _safe_str(val):
            """Convert any value to a PDF-safe string."""
            if val is None:
                return 'N/A'
            return str(val).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        if profile_data:
            # Use correct key: 'basic_data' (not 'basic_info')
            basic_data_pdf = profile_data.get('basic_data', {})
            if basic_data_pdf:
                elements.append(Paragraph("<b>Basic Information</b>", styles['Heading3']))
                info_rows = [
                    ['Field', 'Value'],
                    ['Company Name',    _safe_str(basic_data_pdf.get('name'))],
                    ['HQ Address',      _safe_str(basic_data_pdf.get('hq_address'))],
                    ['CEO',             _safe_str(basic_data_pdf.get('ceo'))],
                    ['Owner / Parent',  _safe_str(basic_data_pdf.get('owner'))],
                    ['Management',      _safe_str(basic_data_pdf.get('management'))],
                    ['Employees (FTE)', _safe_str(basic_data_pdf.get('fte'))],
                    ['Financial Status',_safe_str(basic_data_pdf.get('financials'))],
                    ['Buying Center',   _safe_str(basic_data_pdf.get('buying_center'))],
                    ['Frame Agreements',_safe_str(basic_data_pdf.get('frame_agreements'))],
                ]
                info_table = Table(info_rows, colWidths=[2*inch, 4*inch])
                info_table.setStyle(TableStyle([
                    ('BACKGROUND',   (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
                    ('TEXTCOLOR',    (0, 0), (-1, 0), colors.white),
                    ('ALIGN',        (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME',     (0, 0), (-1,  0), 'Helvetica-Bold'),
                    ('FONTNAME',     (0, 1), ( 0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE',     (0, 0), (-1, -1), 10),
                    ('ROWBACKGROUNDS',(0, 1), (-1, -1),
                     [colors.HexColor('#f9fafb'), colors.white]),
                    ('GRID',         (0, 0), (-1, -1), 0.5, colors.HexColor('#d1d5db')),
                    ('TOPPADDING',   (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING',(0, 0), (-1, -1), 6),
                ]))
                elements.append(info_table)
                elements.append(Spacer(1, 0.2*inch))

                # Text areas
                for label, key in [
                    ('Ownership History',      'ownership_history'),
                    ('Recent News & Facts',     'recent_facts'),
                    ('Company Focus / Vision',  'company_focus'),
                    ('Embargos / ESG Concerns', 'embargos_esg'),
                ]:
                    val = basic_data_pdf.get(key, '')
                    if val and val != 'N/A':
                        elements.append(Paragraph(f"<b>{label}</b>", styles['Heading3']))
                        elements.append(Paragraph(_safe_str(val), styles['CustomBody']))
                        elements.append(Spacer(1, 0.1*inch))

            # History
            history_pdf = profile_data.get('history', {})
            if history_pdf:
                elements.append(Paragraph("<b>Project History & Relationship</b>", styles['Heading3']))
                for label, key in [
                    ('Latest Projects',  'latest_projects'),
                    ('Realized Projects','realized_projects'),
                    ('CRM Rating',       'crm_rating'),
                    ('SMS Relationship', 'sms_relationship'),
                    ('Key Contact',      'key_person'),
                    ('Latest Visits',    'latest_visits'),
                ]:
                    val = history_pdf.get(key, '')
                    if val and val != 'N/A':
                        elements.append(Paragraph(f"<b>{label}:</b> {_safe_str(val)}", styles['CustomBody']))
                elements.append(Spacer(1, 0.1*inch))

            # Context
            context_pdf = profile_data.get('context', {})
            if context_pdf:
                elements.append(Paragraph("<b>Market Context</b>", styles['Heading3']))
                for label, key in [('End Customer', 'end_customer'), ('Market Position', 'market_position')]:
                    val = context_pdf.get(key, '')
                    if val and val != 'N/A':
                        elements.append(Paragraph(f"<b>{label}:</b> {_safe_str(val)}", styles['CustomBody']))
                elements.append(Spacer(1, 0.1*inch))

            # Sales Strategy
            strat_pdf = profile_data.get('sales_strategy', {})
            if strat_pdf:
                elements.append(Paragraph("<b>Strategic Sales Pitch</b>", styles['Heading3']))
                for label, key in [
                    ('Recommended Portfolio', 'recommended_portfolio'),
                    ('Value Proposition',     'value_proposition'),
                    ('Competitive Landscape', 'competitive_landscape'),
                    ('Suggested Next Steps',  'suggested_next_steps'),
                ]:
                    val = strat_pdf.get(key, '')
                    if val and val != 'N/A':
                        elements.append(Paragraph(f"<b>{label}:</b> {_safe_str(val)}", styles['CustomBody']))
                elements.append(Spacer(1, 0.1*inch))

        elements.append(PageBreak())

        # Section 2: Deep Dive Analytics
        if customer_data:
            elements.append(Paragraph("<b>2. Deep Dive Analytics</b>", styles['CustomHeading']))

            projects_list = customer_data.get('projects', [])
            installed_base = customer_data.get('installed_base', [])

            metrics_data = [
                ['Metric', 'Value'],
                ['Total Projects',  str(len(projects_list))],
                ['Total Equipment', str(len(installed_base))],
                ['Total Revenue',   f"${sum([p.get('value', 0) for p in projects_list]):,.0f}"],
            ]

            metrics_table = Table(metrics_data, colWidths=[3*inch, 3*inch])
            metrics_table.setStyle(TableStyle([
                ('BACKGROUND',    (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
                ('TEXTCOLOR',     (0, 0), (-1, 0), colors.white),
                ('ALIGN',         (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME',      (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE',      (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND',    (0, 1), (-1, -1), colors.white),
                ('GRID',          (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
            ]))
            elements.append(metrics_table)
            elements.append(Spacer(1, 0.3*inch))

        # Section 3: Market Intelligence
        # market_intel values are plain strings (not nested dicts)
        if market_intel and isinstance(market_intel, dict):
            elements.append(PageBreak())
            elements.append(Paragraph("<b>3. Market Intelligence</b>", styles['CustomHeading']))

            for mi_label, mi_key in [
                ('Financial Health',      'financial_health'),
                ('Recent Developments',   'recent_developments'),
                ('Market Position',       'market_position'),
                ('Strategic Outlook',     'strategic_outlook'),
                ('Risk Assessment',       'risk_assessment'),
                ('Market Size',           'market_size'),
                ('Growth Trends',         'growth_trends'),
            ]:
                mi_val = market_intel.get(mi_key, '')
                # Value may be a string or a dict with 'summary'
                if isinstance(mi_val, dict):
                    mi_val = mi_val.get('summary', '') or mi_val.get('text', '')
                if mi_val:
                    elements.append(Paragraph(f"<b>{mi_label}</b>", styles['Heading3']))
                    elements.append(Paragraph(_safe_str(mi_val), styles['CustomBody']))
                    elements.append(Spacer(1, 0.1*inch))
        
        # Section 4: Projects
        if projects and len(projects) > 0:
            elements.append(PageBreak())
            elements.append(Paragraph("<b>4. Project Analysis</b>", styles['CustomHeading']))
            
            for idx, project in enumerate(projects[:10], 1):  # Limit to top 10
                project_name = project.get('name', f'Project {idx}')
                elements.append(Paragraph(f"<b>{idx}. {project_name}</b>", styles['Heading3']))
                
                project_data = [
                    ['Status', project.get('status', 'N/A')],
                    ['Value', f"${project.get('value', 0):,.0f}"],
                    ['Budget', f"${project.get('budget', 0):,.0f}"],
                    ['Progress', f"{project.get('progress', 0)}%"],
                ]
                
                for field, value in project_data:
                    elements.append(Paragraph(f"<b>{field}:</b> {value}", styles['CustomBody']))
                
                elements.append(Spacer(1, 0.15*inch))
        
        # Add charts if available
        if charts and PLOTLY_AVAILABLE:
            elements.append(PageBreak())
            elements.append(Paragraph("<b>5. Visualizations</b>", styles['CustomHeading']))
            
            for chart_name, fig in charts.items():
                try:
                    # Convert plotly to image
                    img_bytes = pio.to_image(fig, format='png', width=800, height=400)
                    img_stream = BytesIO(img_bytes)
                    
                    # Add to PDF
                    img = RLImage(img_stream, width=6*inch, height=3*inch)
                    elements.append(img)
                    elements.append(Paragraph(f"<i>{chart_name}</i>", styles['Normal']))
                    elements.append(Spacer(1, 0.2*inch))
                except Exception as e:
                    logger.error(f"Failed to add chart {chart_name} to PDF: {e}")
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer


# Singleton instance
enhanced_export_service = EnhancedExportService()
